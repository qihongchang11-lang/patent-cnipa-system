"""
DOCX exporter for generated patent documents.

Keeps python-docx dependency out of core models.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from core.patent_document import PatentDocument


def _set_default_font(doc: Document, font_name: str = "宋体", font_size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_paragraph_if(doc: Document, text: Optional[str]) -> None:
    if text:
        doc.add_paragraph(text.strip())


def export_patent_docx(patent_doc: PatentDocument, out_path: Path) -> Path:
    """
    Export a single deliverable docx containing: abstract/specification/claims/disclosure.
    Returns out_path for convenience.
    """
    doc = Document()
    _set_default_font(doc)

    doc.add_heading(patent_doc.metadata.title, level=0)

    if patent_doc.abstract:
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(f"技术领域：{patent_doc.abstract.technical_field}")
        doc.add_paragraph(f"发明名称：{patent_doc.abstract.title}")
        _add_paragraph_if(doc, patent_doc.abstract.summary)
        if patent_doc.abstract.main_figure_description:
            doc.add_paragraph(f"主要附图：{patent_doc.abstract.main_figure_description}")

    if patent_doc.specification:
        doc.add_heading("说明书", level=1)
        doc.add_heading("技术领域", level=2)
        _add_paragraph_if(doc, patent_doc.specification.technical_field)
        doc.add_heading("背景技术", level=2)
        _add_paragraph_if(doc, patent_doc.specification.background_art)
        doc.add_heading("发明内容", level=2)
        _add_paragraph_if(doc, patent_doc.specification.invention_content)
        if patent_doc.specification.description_of_drawings:
            doc.add_heading("附图说明", level=2)
            _add_paragraph_if(doc, patent_doc.specification.description_of_drawings)
        doc.add_heading("具体实施方式", level=2)
        _add_paragraph_if(doc, patent_doc.specification.embodiments)

    if patent_doc.claims:
        doc.add_heading("权利要求书", level=1)
        for claim in patent_doc.claims.independent_claims:
            doc.add_paragraph(
                f"{claim.claim_number}. {claim.preamble} {claim.transition} {claim.body}".strip()
            )
        for claim in patent_doc.claims.dependent_claims:
            doc.add_paragraph(
                f"{claim.claim_number}. 根据权利要求{claim.parent_claim}所述的{claim.additional_features}".strip()
            )

    if patent_doc.disclosure:
        doc.add_heading("交底书（具体实施方式）", level=1)
        _add_paragraph_if(doc, patent_doc.disclosure.detailed_description)
        if patent_doc.disclosure.examples:
            doc.add_heading("实施例", level=2)
            for idx, example in enumerate(patent_doc.disclosure.examples, 1):
                doc.add_paragraph(f"实施例 {idx}：")
                _add_paragraph_if(doc, example)
        if patent_doc.disclosure.drawings:
            doc.add_heading("附图", level=2)
            for drawing in patent_doc.disclosure.drawings:
                doc.add_paragraph(f"- {drawing}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path

